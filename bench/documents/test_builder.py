#!/usr/bin/env python3
"""Tests for the deterministic DOCX builder."""

from __future__ import annotations

import contextlib
import importlib.util
import io
import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
BUILDER_PATH = ROOT / "skill/docx/scripts/build_document.py"
MODULE_SPEC = importlib.util.spec_from_file_location("docx_builder", BUILDER_PATH)
assert MODULE_SPEC and MODULE_SPEC.loader
BUILDER = importlib.util.module_from_spec(MODULE_SPEC)
MODULE_SPEC.loader.exec_module(BUILDER)

RU_TITLE = "\u041f\u043b\u0430\u043d \u0437\u0430\u043f\u0443\u0441\u043a\u0430"
RU_SUMMARY = "\u041f\u0438\u043b\u043e\u0442 \u0434\u043e\u0441\u0442\u0438\u0433 \u0446\u0435\u043b\u0435\u0432\u044b\u0445 \u043f\u043e\u043a\u0430\u0437\u0430\u0442\u0435\u043b\u0435\u0439."


def _job(filename: str = "document.docx") -> dict:
    return {
        "format": "docx",
        "audience": "Company leadership",
        "goal": "Approve the pilot",
        "sourceFileIds": ["notes.pdf"],
        "immutableElements": ["Reported actuals"],
        "locale": "ru-RU",
        "filename": filename,
        "acceptanceCriteria": ["The document opens", "Cyrillic renders correctly"],
    }


def _new_spec(filename: str = "document.docx") -> dict:
    return {
        "job": _job(filename),
        "documentType": "memo",
        "title": RU_TITLE,
        "subtitle": "Executive brief",
        "metadata": [{"label": "Status", "value": "Decision required"}],
        "sections": [
            {
                "heading": "Recommendation",
                "blocks": [
                    {"type": "callout", "label": "Decision", "text": RU_SUMMARY},
                    {"type": "bullets", "items": ["Assign an owner", "Define the metric"]},
                    {"type": "numbered", "items": ["Prepare", "Launch", "Review"]},
                    {
                        "type": "table",
                        "columns": ["Metric", "Target", "Owner"],
                        "rows": [["Response", "2 hours", "Operations"]],
                        "widths": [0.4, 0.25, 0.35],
                        "source": "Source: pilot plan, page 4",
                    },
                ],
            }
        ],
        "assumptions": ["The team remains unchanged"],
        "sources": [{"label": "Pilot plan", "location": "notes.pdf, page 4"}],
        "changeLog": [{"target": "Document", "summary": "Created a decision memo"}],
        "repairIterations": 0,
        "outputPdf": True,
    }


def _build_without_render(spec: dict, output: Path):
    BUILDER._job(spec, output)
    source_value = spec.get("inputPath") or spec.get("templatePath")
    source = Path(source_value) if source_value else None
    source_hash = BUILDER._sha256(source) if source else None
    _document, changes, requested = BUILDER._build(spec, output)
    checks, issues = BUILDER._check_structure(output, spec, requested)
    if source:
        immutable = source_hash == BUILDER._sha256(source)
        checks.append(
            {
                "name": "immutable-input",
                "status": "passed" if immutable else "failed",
                "message": "Input remains immutable",
            }
        )
        if not immutable:
            issues.append(BUILDER._issue("input-modified", "critical", "Input changed"))
    report = BUILDER._report(spec, checks, issues, changes, None)
    return report


class ArtifactJobTests(unittest.TestCase):
    def test_pdf_delivery_is_opt_in(self):
        self.assertFalse(BUILDER._output_pdf_requested({}))
        self.assertFalse(BUILDER._output_pdf_requested({"outputPdf": False}))
        self.assertTrue(BUILDER._output_pdf_requested({"outputPdf": True}))
        with self.assertRaisesRegex(ValueError, "outputPdf"):
            BUILDER._output_pdf_requested({"outputPdf": "true"})

    def test_cli_keeps_the_qa_pdf_internal_by_default(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            output = root / "document.docx"
            spec = _new_spec()
            del spec["outputPdf"]
            spec_path = root / "spec.json"
            spec_path.write_text(json.dumps(spec), encoding="utf-8")
            render_check = {"name": "render", "status": "passed", "message": "Rendered"}
            stdout = io.StringIO()

            with mock.patch.object(BUILDER, "_render", return_value=([render_check], [], None)) as render:
                with mock.patch.object(BUILDER.sys, "argv", ["build_document.py", str(spec_path), str(output)]):
                    with contextlib.redirect_stdout(stdout):
                        exit_code = BUILDER.main()

            payload = json.loads(stdout.getvalue())
            render.assert_called_once_with(output, keep_pdf=False)
            self.assertEqual(exit_code, 0)
            self.assertFalse(output.with_suffix(".pdf").exists())
            self.assertEqual(payload["reports"], [f"{output}.artifact-report.json"])

    def test_job_requires_all_shared_contract_fields(self):
        spec = _new_spec()
        del spec["job"]["audience"]
        with tempfile.TemporaryDirectory() as folder:
            with self.assertRaisesRegex(ValueError, "audience"):
                BUILDER._job(spec, Path(folder) / "document.docx")

    def test_job_binds_filename_to_output(self):
        with tempfile.TemporaryDirectory() as folder:
            with self.assertRaisesRegex(ValueError, "match"):
                BUILDER._job(_new_spec("expected.docx"), Path(folder) / "other.docx")

    def test_job_caps_repair_iterations(self):
        spec = _new_spec()
        spec["repairIterations"] = 3
        with tempfile.TemporaryDirectory() as folder:
            with self.assertRaisesRegex(ValueError, "between 0 and 2"):
                BUILDER._job(spec, Path(folder) / "document.docx")

    def test_job_refuses_in_place_edits(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "source.docx"
            Document().save(source)
            spec = {
                "job": _job("source.docx"),
                "inputPath": str(source),
                "edits": [{"oldText": "old", "newText": "new"}],
            }
            with self.assertRaisesRegex(ValueError, "must differ"):
                BUILDER._job(spec, source)


class NewDocumentTests(unittest.TestCase):
    def test_source_urls_reject_unsafe_relationship_targets(self):
        unsafe_urls = [
            "file:///Users/example/secret.txt",
            "smb://files.example.org/share",
            "javascript:alert(1)",
            "https:///missing-host",
            "https://user:password@example.org/source",
            "https://example.org/source\nfile:///tmp/secret",
            "https://example.org/source\x00hidden",
            " https://example.org/source",
            "https://example.org/source\u0085hidden",
            "https://example.org:invalid/source",
            "https://[::1/source",
        ]

        with tempfile.TemporaryDirectory() as folder:
            for index, url in enumerate(unsafe_urls):
                with self.subTest(url=url):
                    output = Path(folder) / f"document-{index}.docx"
                    spec = _new_spec(output.name)
                    spec["sources"] = [{"label": "Source", "url": url}]

                    with self.assertRaisesRegex(ValueError, "sources\\[\\]\\.url"):
                        BUILDER._job(spec, output)
                    self.assertFalse(output.exists())

    def test_web_sources_create_external_http_relationships(self):
        source_urls = [
            "http://example.org/research/service-methodology",
            "https://example.org/research/pilot-evaluation?version=2",
        ]
        spec = _new_spec()
        spec["sources"] = [
            {"label": f"Methodology {index}", "url": url}
            for index, url in enumerate(source_urls, start=1)
        ]

        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "document.docx"
            report = _build_without_render(spec, output)
            with zipfile.ZipFile(output) as package:
                relationships = package.read("word/_rels/document.xml.rels").decode("utf-8")

        self.assertEqual(report["status"], "ready", report["issues"])
        for source_url in source_urls:
            self.assertEqual(relationships.count(source_url.replace("&", "&amp;")), 1)
        self.assertEqual(relationships.count('TargetMode="External"'), len(source_urls))

    def test_new_document_uses_semantic_styles_lists_and_fixed_tables(self):
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "document.docx"
            report = _build_without_render(_new_spec(), output)
            document = Document(output)

            self.assertEqual(report["status"], "ready", report["issues"])
            self.assertEqual(round(document.sections[0].page_width.cm, 1), 21.0)
            self.assertEqual(round(document.sections[0].page_height.cm, 1), 29.7)
            self.assertTrue(any(p.style.name == "Heading 1" for p in document.paragraphs))
            numbered = [
                p
                for p in document.paragraphs
                if p._p.get_or_add_pPr().find(qn("w:numPr")) is not None
            ]
            self.assertGreaterEqual(len(numbered), 6)
            self.assertFalse(BUILDER._table_geometry_issues(document))
            self.assertIn(RU_TITLE, "\n".join(p.text for p in document.paragraphs))

    def test_independent_lists_restart_with_distinct_numbering_ids(self):
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "document.docx"
            _build_without_render(_new_spec(), output)
            document = Document(output)

            numbering_ids = {}
            for paragraph in document.paragraphs:
                num_pr = paragraph._p.get_or_add_pPr().find(qn("w:numPr"))
                if num_pr is None:
                    continue
                num_id = num_pr.find(qn("w:numId"))
                if num_id is not None:
                    numbering_ids[paragraph.text] = num_id.get(qn("w:val"))

            independent_starts = [
                numbering_ids["Assign an owner"],
                numbering_ids["Prepare"],
                numbering_ids["The team remains unchanged"],
                numbering_ids["Pilot plan — notes.pdf, page 4"],
            ]
            self.assertEqual(len(set(independent_starts)), len(independent_starts))

    def test_homepage_only_source_fails_traceability(self):
        spec = _new_spec()
        spec["sources"] = [{"label": "Publisher", "url": "https://example.com/"}]
        with tempfile.TemporaryDirectory() as folder:
            report = _build_without_render(spec, Path(folder) / "document.docx")
        self.assertEqual(report["status"], "needs_review")
        self.assertTrue(any(issue["code"] == "unspecific-source" for issue in report["issues"]))

    def test_pdf_report_inherits_verified_metadata(self):
        spec = _new_spec()
        checks = [{"name": "reopen", "status": "passed", "message": "ok"}]
        report = BUILDER._report(
            spec,
            checks,
            [],
            [{"target": "Document", "summary": "Created"}],
            Path("document.pdf"),
        )
        self.assertEqual(report["format"], "docx")
        self.assertEqual(report["previewAssets"], [{"filename": "document.pdf", "kind": "pdf"}])
        self.assertEqual(report["repairIterations"], 0)

    def test_russian_default_change_log_is_localized(self):
        spec = _new_spec()
        spec["changeLog"] = []
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "document.docx"
            _document, changes, _requested = BUILDER._build(spec, output)

        self.assertTrue(any("\u0400" <= char <= "\u04ff" for char in changes[0]["summary"]))


class TemplateAndEditTests(unittest.TestCase):
    def test_template_placeholder_can_span_runs_without_modifying_source(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "template.docx"
            output = root / "filled.docx"
            document = Document()
            paragraph = document.add_paragraph(style="Title")
            first = paragraph.add_run("{{REPORT_")
            first.bold = True
            paragraph.add_run("TITLE}}")
            document.sections[0].top_margin = BUILDER.Cm(1.7)
            document.save(source)
            source_hash = BUILDER._sha256(source)
            spec = {
                "job": {**_job("filled.docx"), "templateFileId": "template.docx"},
                "templatePath": str(source),
                "placeholders": {"{{REPORT_TITLE}}": RU_TITLE},
                "outputPdf": False,
            }

            report = _build_without_render(spec, output)
            filled = Document(output)

            self.assertEqual(report["status"], "ready", report["issues"])
            self.assertEqual(BUILDER._sha256(source), source_hash)
            self.assertEqual(filled.paragraphs[0].text, RU_TITLE)
            self.assertTrue(filled.paragraphs[0].runs[0].bold)
            self.assertEqual(filled.sections[0].top_margin, Document(source).sections[0].top_margin)

    def test_template_replaces_placeholders_in_tables_headers_and_footers(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "template.docx"
            output = root / "filled.docx"
            document = Document()
            document.add_table(rows=1, cols=1).cell(0, 0).text = "{{TABLE_VALUE}}"
            document.sections[0].header.paragraphs[0].text = "{{HEADER_VALUE}}"
            document.sections[0].footer.paragraphs[0].text = "{{FOOTER_VALUE}}"
            document.save(source)
            source_hash = BUILDER._sha256(source)
            spec = {
                "job": {**_job("filled.docx"), "templateFileId": "template.docx"},
                "templatePath": str(source),
                "placeholders": {
                    "{{TABLE_VALUE}}": "Approved value",
                    "{{HEADER_VALUE}}": "Approved header",
                    "{{FOOTER_VALUE}}": "Controlled copy",
                },
                "outputPdf": False,
            }

            report = _build_without_render(spec, output)
            filled = Document(output)

            self.assertEqual(report["status"], "ready", report["issues"])
            self.assertEqual(BUILDER._sha256(source), source_hash)
            self.assertEqual(filled.tables[0].cell(0, 0).text, "Approved value")
            self.assertEqual(filled.sections[0].header.paragraphs[0].text, "Approved header")
            self.assertEqual(filled.sections[0].footer.paragraphs[0].text, "Controlled copy")

    def test_template_with_unresolved_placeholder_needs_review(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "template.docx"
            output = root / "filled.docx"
            document = Document()
            document.add_paragraph("{{TITLE}} / {{OWNER}}")
            document.save(source)
            spec = {
                "job": {**_job("filled.docx"), "templateFileId": "template.docx"},
                "templatePath": str(source),
                "placeholders": {"{{TITLE}}": "Pilot report"},
            }
            report = _build_without_render(spec, output)
        self.assertEqual(report["status"], "needs_review")
        self.assertTrue(any(issue["code"] == "unresolved-placeholder" for issue in report["issues"]))

    def test_targeted_edit_changes_only_requested_word_part(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.docx"
            output = root / "updated.docx"
            document = Document()
            document.add_heading("Pilot", level=1)
            document.add_paragraph("Duration: four weeks.")
            document.sections[0].header.paragraphs[0].text = "Stable header"
            document.save(source)
            source_hash = BUILDER._sha256(source)
            spec = {
                "job": _job("updated.docx"),
                "inputPath": str(source),
                "edits": [
                    {
                        "oldText": "four weeks",
                        "newText": "six weeks",
                        "summary": "Updated the approved duration",
                    }
                ],
                "outputPdf": False,
            }

            report = _build_without_render(spec, output)

            self.assertEqual(report["status"], "ready", report["issues"])
            self.assertEqual(BUILDER._sha256(source), source_hash)
            self.assertIn("six weeks", "\n".join(p.text for p in Document(output).paragraphs))
            scope = next(check for check in report["qaChecks"] if check["name"] == "targeted-edit-scope")
            self.assertEqual(scope["status"], "passed")

    def test_targeted_edit_rejects_ambiguous_text(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.docx"
            output = root / "updated.docx"
            document = Document()
            document.add_paragraph("old value")
            document.add_paragraph("old value")
            document.save(source)
            spec = {
                "job": _job("updated.docx"),
                "inputPath": str(source),
                "edits": [{"oldText": "old value", "newText": "new value"}],
            }
            BUILDER._job(spec, output)
            with self.assertRaisesRegex(ValueError, "exactly once"):
                BUILDER._build(spec, output)


class RenderIntegrationTests(unittest.TestCase):
    @unittest.skipUnless(BUILDER.shutil.which("soffice"), "LibreOffice is unavailable")
    def test_real_render_produces_pdf_and_cyrillic_check(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            output = root / "document.docx"
            spec = _new_spec()
            BUILDER._job(spec, output)
            BUILDER._build(spec, output)
            checks, issues, pdf = BUILDER._render(output, keep_pdf=True)

            self.assertIsNotNone(pdf)
            self.assertTrue(pdf and pdf.is_file())
            self.assertFalse([issue for issue in issues if issue["severity"] == "critical"], issues)
            self.assertEqual(next(c for c in checks if c["name"] == "render")["status"], "passed")
            self.assertEqual(next(c for c in checks if c["name"] == "cyrillic-render")["status"], "passed")


if __name__ == "__main__":
    unittest.main()
