#!/usr/bin/env python3
"""Run the Russian-first DOCX golden matrix in the real office toolchain."""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageOps


ROOT = Path(__file__).resolve().parents[2]
BENCH_DIR = Path(__file__).resolve().parent
BUILDER = ROOT / "skill/docx/scripts/build_document.py"
CASES_PATH = BENCH_DIR / "cases.json"
RESULTS_DIR = BENCH_DIR / "results"
RENDER_EVIDENCE_ALGORITHM = "sha256-json-v1"

RU = {
    "title": "\u041f\u043b\u0430\u043d \u0437\u0430\u043f\u0443\u0441\u043a\u0430 \u043f\u0438\u043b\u043e\u0442\u0430",
    "summary": "\u041f\u0438\u043b\u043e\u0442 \u0441\u043d\u0438\u0436\u0430\u0435\u0442 \u043e\u043f\u0435\u0440\u0430\u0446\u0438\u043e\u043d\u043d\u044b\u0439 \u0440\u0438\u0441\u043a \u0438 \u0434\u0430\u0451\u0442 \u0438\u0437\u043c\u0435\u0440\u0438\u043c\u044b\u0439 \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442.",
    "recommendation": "\u0420\u0435\u043a\u043e\u043c\u0435\u043d\u0434\u0430\u0446\u0438\u044f",
    "analysis": "\u0410\u043d\u0430\u043b\u0438\u0437 \u0438 \u043e\u0431\u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0435",
    "next": "\u0421\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0435 \u0448\u0430\u0433\u0438",
    "decision": "\u0420\u0435\u0448\u0435\u043d\u0438\u0435",
    "owner": "\u0412\u043b\u0430\u0434\u0435\u043b\u0435\u0446",
    "metric": "\u041c\u0435\u0442\u0440\u0438\u043a\u0430",
    "target": "\u0426\u0435\u043b\u044c",
    "status": "\u0421\u0442\u0430\u0442\u0443\u0441",
    "subtitle": "\u0420\u0430\u0431\u043e\u0447\u0438\u0439 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442 \u0434\u043b\u044f \u0440\u0443\u043a\u043e\u0432\u043e\u0434\u0441\u0442\u0432\u0430 | \u0430\u0432\u0433\u0443\u0441\u0442 2026",
    "audience": "\u0420\u0443\u043a\u043e\u0432\u043e\u0434\u0441\u0442\u0432\u043e \u043a\u043e\u043c\u043f\u0430\u043d\u0438\u0438",
    "audience_label": "\u0410\u0443\u0434\u0438\u0442\u043e\u0440\u0438\u044f",
    "decision_required": "\u0422\u0440\u0435\u0431\u0443\u0435\u0442\u0441\u044f \u0440\u0435\u0448\u0435\u043d\u0438\u0435",
    "assign_owner": "\u041d\u0430\u0437\u043d\u0430\u0447\u0438\u0442\u044c \u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u0432\u043b\u0430\u0434\u0435\u043b\u044c\u0446\u0430",
    "approve_measurement": "\u0423\u0442\u0432\u0435\u0440\u0434\u0438\u0442\u044c \u043c\u0435\u0442\u043e\u0434\u0438\u043a\u0443 \u0438\u0437\u043c\u0435\u0440\u0435\u043d\u0438\u044f",
    "review_six_weeks": "\u041f\u043e\u0434\u0432\u0435\u0441\u0442\u0438 \u0438\u0442\u043e\u0433\u0438 \u0447\u0435\u0440\u0435\u0437 \u0448\u0435\u0441\u0442\u044c \u043d\u0435\u0434\u0435\u043b\u044c",
    "response_time": "\u0412\u0440\u0435\u043c\u044f \u043e\u0442\u0432\u0435\u0442\u0430",
    "under_two_hours": "\u043c\u0435\u043d\u0435\u0435 2 \u0447\u0430\u0441\u043e\u0432",
    "operations": "\u041e\u043f\u0435\u0440\u0430\u0446\u0438\u043e\u043d\u043d\u0430\u044f \u043a\u043e\u043c\u0430\u043d\u0434\u0430",
    "satisfaction": "\u0423\u0434\u043e\u0432\u043b\u0435\u0442\u0432\u043e\u0440\u0451\u043d\u043d\u043e\u0441\u0442\u044c",
    "at_least_85": "\u043d\u0435 \u043c\u0435\u043d\u0435\u0435 85%",
    "service_lead": "\u0420\u0443\u043a\u043e\u0432\u043e\u0434\u0438\u0442\u0435\u043b\u044c \u0441\u0435\u0440\u0432\u0438\u0441\u0430",
    "prepare_data": "\u041f\u043e\u0434\u0433\u043e\u0442\u043e\u0432\u0438\u0442\u044c \u0438\u0441\u0445\u043e\u0434\u043d\u044b\u0435 \u0434\u0430\u043d\u043d\u044b\u0435",
    "run_pilot": "\u041f\u0440\u043e\u0432\u0435\u0441\u0442\u0438 \u043a\u043e\u043d\u0442\u0440\u043e\u043b\u0438\u0440\u0443\u0435\u043c\u044b\u0439 \u043f\u0438\u043b\u043e\u0442",
    "review_metrics": "\u041f\u0440\u043e\u0432\u0435\u0440\u0438\u0442\u044c \u0446\u0435\u043b\u0435\u0432\u044b\u0435 \u043c\u0435\u0442\u0440\u0438\u043a\u0438",
    "decide_scale": "\u041f\u0440\u0438\u043d\u044f\u0442\u044c \u0440\u0435\u0448\u0435\u043d\u0438\u0435 \u043e \u043c\u0430\u0441\u0448\u0442\u0430\u0431\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u0438",
    "team_assumption": "\u0421\u043e\u0441\u0442\u0430\u0432 \u043f\u0438\u043b\u043e\u0442\u043d\u043e\u0439 \u043a\u043e\u043c\u0430\u043d\u0434\u044b \u043d\u0435 \u043c\u0435\u043d\u044f\u0435\u0442\u0441\u044f \u0432 \u0442\u0435\u0447\u0435\u043d\u0438\u0435 \u0448\u0435\u0441\u0442\u0438 \u043d\u0435\u0434\u0435\u043b\u044c",
    "project_notes": "\u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044b \u043f\u0440\u043e\u0435\u043a\u0442\u0430",
    "operating_metrics": "\u041e\u043f\u0435\u0440\u0430\u0446\u0438\u043e\u043d\u043d\u044b\u0435 \u043c\u0435\u0442\u0440\u0438\u043a\u0438",
    "purpose_scope": "\u0426\u0435\u043b\u044c \u0438 \u043e\u0431\u043b\u0430\u0441\u0442\u044c \u043f\u0440\u0438\u043c\u0435\u043d\u0435\u043d\u0438\u044f",
    "roles": "\u0420\u043e\u043b\u0438",
    "role": "\u0420\u043e\u043b\u044c",
    "responsibility": "\u041e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0441\u0442\u044c",
    "escalation": "\u042d\u0441\u043a\u0430\u043b\u0430\u0446\u0438\u044f",
    "procedure": "\u041f\u043e\u0440\u044f\u0434\u043e\u043a \u0434\u0435\u0439\u0441\u0442\u0432\u0438\u0439",
    "controls": "\u041a\u043e\u043d\u0442\u0440\u043e\u043b\u044c\u043d\u044b\u0435 \u043c\u0435\u0440\u044b",
    "duration_four": "\u0423\u0442\u0432\u0435\u0440\u0436\u0434\u0451\u043d\u043d\u0430\u044f \u0434\u043b\u0438\u0442\u0435\u043b\u044c\u043d\u043e\u0441\u0442\u044c \u043f\u0438\u043b\u043e\u0442\u0430: \u0447\u0435\u0442\u044b\u0440\u0435 \u043d\u0435\u0434\u0435\u043b\u0438.",
    "duration_four_words": "\u0447\u0435\u0442\u044b\u0440\u0435 \u043d\u0435\u0434\u0435\u043b\u0438",
    "duration_six": "\u0448\u0435\u0441\u0442\u044c \u043d\u0435\u0434\u0435\u043b\u044c",
    "baseline_excludes": "\u0411\u0430\u0437\u043e\u0432\u044b\u0439 \u0440\u0430\u0441\u0447\u0451\u0442 \u043d\u0435 \u0432\u043a\u043b\u044e\u0447\u0430\u0435\u0442 \u0440\u0430\u0437\u043e\u0432\u044b\u0435 \u0440\u0430\u0431\u043e\u0442\u044b \u043f\u043e \u043c\u0438\u0433\u0440\u0430\u0446\u0438\u0438",
    "exchange_rate": "\u041a\u0443\u0440\u0441 \u0432\u0430\u043b\u044e\u0442 \u0437\u0430\u0444\u0438\u043a\u0441\u0438\u0440\u043e\u0432\u0430\u043d \u043d\u0430 \u043f\u043b\u0430\u043d\u043e\u0432\u043e\u043c \u0443\u0440\u043e\u0432\u043d\u0435",
    "research_page": "\u041c\u0435\u0442\u043e\u0434\u0438\u043a\u0430 \u043e\u0446\u0435\u043d\u043a\u0438 \u043f\u0438\u043b\u043e\u0442\u0430",
    "source_note": "\u0418\u0441\u0442\u043e\u0447\u043d\u0438\u043a: project-notes.pdf, \u0441\u0442\u0440. 4",
    "source_metrics": "\u0418\u0441\u0442\u043e\u0447\u043d\u0438\u043a: metrics.xlsx, \u043b\u0438\u0441\u0442 Baseline",
    "section": "\u0420\u0430\u0437\u0434\u0435\u043b",
    "role_owner": "\u0412\u043b\u0430\u0434\u0435\u043b\u0435\u0446 \u043f\u0440\u043e\u0446\u0435\u0441\u0441\u0430",
    "role_owner_task": "\u0423\u0442\u0432\u0435\u0440\u0436\u0434\u0430\u0435\u0442 \u0437\u0430\u043f\u0443\u0441\u043a \u0438 \u0438\u0442\u043e\u0433",
    "role_owner_escalation": "\u041a\u0443\u0440\u0430\u0442\u043e\u0440 \u043f\u0440\u043e\u0435\u043a\u0442\u0430",
    "role_operator": "\u0418\u0441\u043f\u043e\u043b\u043d\u0438\u0442\u0435\u043b\u044c",
    "role_operator_task": "\u0412\u044b\u043f\u043e\u043b\u043d\u044f\u0435\u0442 \u0448\u0430\u0433\u0438 \u043f\u0440\u043e\u0446\u0435\u0434\u0443\u0440\u044b",
    "role_reviewer": "\u041f\u0440\u043e\u0432\u0435\u0440\u044f\u044e\u0449\u0438\u0439",
    "role_reviewer_task": "\u041f\u0440\u043e\u0432\u0435\u0440\u044f\u0435\u0442 \u043f\u043e\u0434\u0442\u0432\u0435\u0440\u0436\u0434\u0435\u043d\u0438\u044f",
    "role_reviewer_escalation": "\u041e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0435\u043d\u043d\u044b\u0439 \u0437\u0430 \u043a\u043e\u043d\u0442\u0440\u043e\u043b\u044c",
}

RU_SOP_STEPS = [
    "\u041f\u0440\u043e\u0432\u0435\u0440\u0438\u0442\u044c \u043f\u043e\u043b\u043d\u043e\u0442\u0443 \u0432\u0445\u043e\u0434\u043d\u044b\u0445 \u0434\u0430\u043d\u043d\u044b\u0445",
    "\u041d\u0430\u0437\u043d\u0430\u0447\u0438\u0442\u044c \u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u0437\u0430 \u0437\u0430\u043f\u0443\u0441\u043a",
    "\u0417\u0430\u0444\u0438\u043a\u0441\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u0431\u0430\u0437\u043e\u0432\u044b\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u044f \u043c\u0435\u0442\u0440\u0438\u043a",
    "\u0421\u043e\u0433\u043b\u0430\u0441\u043e\u0432\u0430\u0442\u044c \u043a\u0440\u0438\u0442\u0435\u0440\u0438\u0438 \u043e\u0441\u0442\u0430\u043d\u043e\u0432\u043a\u0438",
    "\u041f\u043e\u0434\u0433\u043e\u0442\u043e\u0432\u0438\u0442\u044c \u0440\u0430\u0431\u043e\u0447\u0443\u044e \u0441\u0440\u0435\u0434\u0443 \u0438 \u0434\u043e\u0441\u0442\u0443\u043f\u044b",
    "\u041f\u0440\u043e\u0432\u0435\u0441\u0442\u0438 \u043a\u043e\u043d\u0442\u0440\u043e\u043b\u044c\u043d\u044b\u0439 \u0437\u0430\u043f\u0443\u0441\u043a",
    "\u0417\u0430\u0444\u0438\u043a\u0441\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442 \u043a\u0430\u0436\u0434\u043e\u0433\u043e \u044d\u0442\u0430\u043f\u0430",
    "\u0417\u0430\u0440\u0435\u0433\u0438\u0441\u0442\u0440\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u043e\u0442\u043a\u043b\u043e\u043d\u0435\u043d\u0438\u044f \u0438 \u0440\u0435\u0448\u0435\u043d\u0438\u044f",
    "\u0421\u0432\u0435\u0440\u0438\u0442\u044c \u0438\u0442\u043e\u0433\u043e\u0432\u044b\u0435 \u043f\u043e\u043a\u0430\u0437\u0430\u0442\u0435\u043b\u0438 \u0441 \u0431\u0430\u0437\u043e\u0432\u044b\u043c \u0443\u0440\u043e\u0432\u043d\u0435\u043c",
    "\u041f\u043e\u0434\u0433\u043e\u0442\u043e\u0432\u0438\u0442\u044c \u0437\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435 \u0434\u043b\u044f \u0440\u0443\u043a\u043e\u0432\u043e\u0434\u0441\u0442\u0432\u0430",
    "\u041f\u0440\u043e\u0432\u0435\u0441\u0442\u0438 \u043d\u0435\u0437\u0430\u0432\u0438\u0441\u0438\u043c\u0443\u044e \u043f\u0440\u043e\u0432\u0435\u0440\u043a\u0443 \u043f\u043e\u0434\u0442\u0432\u0435\u0440\u0436\u0434\u0435\u043d\u0438\u0439",
    "\u0423\u0442\u0432\u0435\u0440\u0434\u0438\u0442\u044c \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442 \u0438 \u0430\u0440\u0445\u0438\u0432\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u043c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044b",
]

RU_SOP_CONTROLS = [
    "\u041e\u0441\u0442\u0430\u043d\u043e\u0432\u0438\u0442\u044c \u043f\u0440\u043e\u0446\u0435\u0441\u0441, \u0435\u0441\u043b\u0438 \u043e\u0442\u0441\u0443\u0442\u0441\u0442\u0432\u0443\u044e\u0442 \u043e\u0431\u044f\u0437\u0430\u0442\u0435\u043b\u044c\u043d\u044b\u0435 \u043f\u043e\u0434\u0442\u0432\u0435\u0440\u0436\u0434\u0435\u043d\u0438\u044f",
    "\u0417\u0430\u0444\u0438\u043a\u0441\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u0438\u0441\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u044f \u0432 \u0436\u0443\u0440\u043d\u0430\u043b\u0435 \u0438\u0437\u043c\u0435\u043d\u0435\u043d\u0438\u0439",
    "\u041d\u0435\u043c\u0435\u0434\u043b\u0435\u043d\u043d\u043e \u044d\u0441\u043a\u0430\u043b\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u043a\u0440\u0438\u0442\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u043e\u0442\u043a\u043b\u043e\u043d\u0435\u043d\u0438\u044f",
]


def _job(case_id: str, filename: str) -> dict[str, Any]:
    return {
        "format": "docx",
        "audience": "Company leadership",
        "goal": f"Approve the outcome of {case_id}",
        "sourceFileIds": ["project-notes.pdf", "metrics.xlsx"],
        "immutableElements": ["Reported actuals"],
        "locale": "ru-RU",
        "filename": filename,
        "acceptanceCriteria": [
            "The document reopens and remains editable",
            "Cyrillic renders correctly",
            "Sources and assumptions are traceable",
        ],
    }


def _paragraph(text: str, repeat: int = 1) -> dict[str, Any]:
    return {"type": "paragraph", "text": " ".join([text] * repeat)}


def _base_spec(case_id: str, filename: str, document_type: str) -> dict[str, Any]:
    return {
        "job": _job(case_id, filename),
        "documentType": document_type,
        "title": RU["title"],
        "subtitle": RU["subtitle"],
        "metadata": [
            {"label": RU["audience_label"], "value": RU["audience"]},
            {"label": RU["status"], "value": RU["decision_required"]},
        ],
        "sections": [
            {
                "heading": RU["recommendation"],
                "blocks": [
                    {"type": "callout", "label": RU["decision"], "text": RU["summary"]},
                    _paragraph(RU["summary"]),
                    {"type": "bullets", "items": [RU["assign_owner"], RU["approve_measurement"], RU["review_six_weeks"]]},
                ],
            },
            {
                "heading": RU["analysis"],
                "blocks": [
                    _paragraph(RU["summary"], 3),
                    {
                        "type": "table",
                        "columns": [RU["metric"], RU["target"], RU["owner"]],
                        "rows": [[RU["response_time"], RU["under_two_hours"], RU["operations"]], [RU["satisfaction"], RU["at_least_85"], RU["service_lead"]]],
                        "widths": [0.4, 0.25, 0.35],
                        "source": RU["source_note"],
                    },
                ],
            },
            {
                "heading": RU["next"],
                "blocks": [
                    {"type": "numbered", "items": [RU["prepare_data"], RU["run_pilot"], RU["review_metrics"], RU["decide_scale"]]}
                ],
            },
        ],
        "assumptions": [RU["team_assumption"]],
        "sources": [
            {"label": RU["project_notes"], "location": "project-notes.pdf, \u0441\u0442\u0440. 4\u20138"},
            {"label": RU["operating_metrics"], "location": "metrics.xlsx, \u043b\u0438\u0441\u0442 Baseline"},
        ],
        "changeLog": [{"target": "Document", "summary": f"Created golden scenario {case_id}"}],
        "repairIterations": 0,
        "outputPdf": True,
    }


def _make_template(path: Path) -> None:
    document = Document()
    document.sections[0].top_margin = BUILDER_MODULE.Cm(1.6)
    title = document.add_paragraph(style="Title")
    title.add_run("{{REPORT_").bold = True
    title.add_run("TITLE}}")
    document.add_paragraph("{{SUMMARY}}")
    document.sections[0].header.paragraphs[0].text = RU["audience"]
    document.sections[0].footer.paragraphs[0].text = RU["status"]
    document.save(path)


def _make_edit_source(path: Path) -> None:
    document = Document()
    document.add_heading(RU["title"], level=1)
    document.add_paragraph(RU["duration_four"])
    document.add_paragraph(RU["summary"])
    document.sections[0].header.paragraphs[0].text = RU["decision_required"]
    document.save(path)


def _case_spec(case: dict[str, Any], filename: str, run_dir: Path) -> tuple[dict[str, Any], Path | None]:
    case_id = case["id"]
    if case_id == "ru-template-fidelity":
        template = run_dir / "template.docx"
        _make_template(template)
        return (
            {
                "job": {**_job(case_id, filename), "templateFileId": "template.docx"},
                "templatePath": str(template),
                "placeholders": {"{{REPORT_TITLE}}": RU["title"], "{{SUMMARY}}": RU["summary"]},
                "changeLog": [{"target": "Template", "summary": "Filled the approved template"}],
                "repairIterations": 0,
                "outputPdf": True,
            },
            template,
        )
    if case_id == "ru-targeted-edit":
        source = run_dir / "source.docx"
        _make_edit_source(source)
        return (
            {
                "job": _job(case_id, filename),
                "inputPath": str(source),
                "edits": [{"oldText": RU["duration_four_words"], "newText": RU["duration_six"], "summary": "Updated the approved pilot duration"}],
                "repairIterations": 0,
                "outputPdf": True,
            },
            source,
        )

    spec = _base_spec(case_id, filename, case["documentType"])
    if case_id == "ru-business-report":
        for section in spec["sections"]:
            section["blocks"].insert(0, _paragraph(RU["summary"], 5))
    elif case_id == "ru-sop":
        spec["sections"] = [
            {"heading": RU["purpose_scope"], "blocks": [_paragraph(RU["summary"], 2)]},
            {"heading": RU["roles"], "blocks": [{"type": "table", "columns": [RU["role"], RU["responsibility"], RU["escalation"]], "rows": [[RU["role_owner"], RU["role_owner_task"], RU["role_owner_escalation"]], [RU["role_operator"], RU["role_operator_task"], RU["role_owner"]], [RU["role_reviewer"], RU["role_reviewer_task"], RU["role_reviewer_escalation"]]], "widths": [0.22, 0.48, 0.30]}]},
            {"heading": RU["procedure"], "blocks": [{"type": "numbered", "items": RU_SOP_STEPS}, {"type": "page_break"}]},
            {"heading": RU["controls"], "blocks": [{"type": "bullets", "items": RU_SOP_CONTROLS}]},
        ]
    elif case_id == "ru-long-cyrillic":
        spec["sections"][0]["heading"] = RU["recommendation"] + ": " + RU["summary"]
        spec["sections"][1]["blocks"] = [_paragraph(RU["summary"], 18)]
    elif case_id == "ru-multipage-table":
        rows = [[f"{RU['metric']} {index:02d}", f"{RU['target']} {index}", f"{RU['owner']} {index % 5}"] for index in range(1, 76)]
        spec["sections"] = [{"heading": RU["analysis"], "blocks": [{"type": "table", "columns": [RU["metric"], RU["target"], RU["owner"]], "rows": rows, "widths": [0.4, 0.25, 0.35], "source": RU["source_metrics"]}]}]
    elif case_id == "ru-source-traceability":
        spec["assumptions"] = [RU["baseline_excludes"], RU["exchange_rate"]]
        spec["sources"].append({"label": RU["research_page"], "url": "https://example.org/research/pilot-evaluation", "accessed": "2026-08-31"})
    elif case_id == "ru-page-breaks":
        spec["sections"] = [
            {"heading": f"{RU['section']} {index}", "blocks": [_paragraph(RU["summary"], 8), *([{"type": "page_break"}] if index < 3 else [])]}
            for index in range(1, 4)
        ]
    elif case_id == "ru-mixed-typography":
        spec["sections"][0]["blocks"].append(_paragraph("ARR grew by 18.4% to 12.7 million RUB on 31.08.2026; SLA and NPS remain within target."))
    return spec, None


def _reset_run_dir(run_dir: Path) -> None:
    shutil.rmtree(run_dir, ignore_errors=True)
    run_dir.mkdir(parents=True)


def build_render_evidence_digest(
    results: list[dict[str, Any]], expected_cases: list[str]
) -> tuple[str | None, list[str]]:
    evidence: dict[str, list[str]] = {}
    for result in results:
        if result.get("run") != 1 or result.get("case") not in expected_cases:
            continue
        hashes = result.get("imageHashes")
        if isinstance(hashes, list) and hashes and all(isinstance(value, str) and value for value in hashes):
            evidence[str(result["case"])] = hashes
    missing = sorted(case_id for case_id in expected_cases if case_id not in evidence)
    if missing:
        return None, ["render evidence is missing for: " + ", ".join(missing)]
    canonical = json.dumps(
        {case_id: evidence[case_id] for case_id in sorted(expected_cases)},
        separators=(",", ":"),
        ensure_ascii=False,
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest(), []


def evaluate_matrix_coverage(
    results: list[dict[str, Any]], expected_cases: list[str], expected_runs: int
) -> list[str]:
    issues: list[str] = []
    observed = {(str(result.get("case")), int(result.get("run", 0))) for result in results}
    for case_id in expected_cases:
        for run_number in range(1, expected_runs + 1):
            if (case_id, run_number) not in observed:
                issues.append(f"missing result for {case_id} run {run_number}")
    unexpected = sorted(case_id for case_id, _ in observed if case_id not in expected_cases)
    if unexpected:
        issues.append("unexpected cases: " + ", ".join(unexpected))
    for result in results:
        if result.get("status") != "ready":
            issues.append(f"{result.get('case')} run {result.get('run')} is not ready")
        if not isinstance(result.get("pages"), int) or int(result["pages"]) < 1:
            issues.append(f"{result.get('case')} run {result.get('run')} has no rendered pages")
    return issues


def _render_pages(pdf: Path, run_dir: Path) -> list[Path]:
    raster = shutil.which("pdftoppm")
    if not raster:
        raise RuntimeError("pdftoppm is unavailable")
    result = subprocess.run([raster, "-png", "-r", "100", str(pdf), str(run_dir / "page")], capture_output=True, text=True, timeout=120, check=False)
    images = sorted(run_dir.glob("page-*.png"))
    if result.returncode != 0 or not images:
        raise RuntimeError((result.stderr or result.stdout or "PDF rasterization failed")[-400:])
    return images


def _write_montage(images: list[Path], output: Path) -> None:
    thumb_size = (330, 467)
    gap = 24
    label_height = 28
    columns = 3
    rows = (len(images) + columns - 1) // columns
    canvas = Image.new("RGB", (columns * (thumb_size[0] + gap) + gap, rows * (thumb_size[1] + label_height + gap) + gap), "#E8ECF2")
    draw = ImageDraw.Draw(canvas)
    for index, path in enumerate(images):
        with Image.open(path) as source:
            thumb = ImageOps.contain(source.convert("RGB"), thumb_size)
        column = index % columns
        row = index // columns
        x = gap + column * (thumb_size[0] + gap)
        y = gap + row * (thumb_size[1] + label_height + gap)
        canvas.paste(thumb, (x, y))
        draw.text((x, y + thumb_size[1] + 5), f"Page {index + 1}", fill="#20242A")
    canvas.save(output)


def _num_paragraphs(document: Document) -> int:
    return sum(1 for paragraph in document.paragraphs if paragraph._p.get_or_add_pPr().find(qn("w:numPr")) is not None)


def _assert_requirements(case: dict[str, Any], output: Path, report: dict[str, Any], source: Path | None, source_hash: str | None, images: list[Path]) -> list[str]:
    requirements = case["requirements"]
    document = Document(output)
    issues: list[str] = []
    if report.get("status") != "ready":
        issues.append(f"artifact status is {report.get('status')}")
    if len(images) < int(requirements.get("minPages", 1)):
        issues.append(f"rendered {len(images)} pages, expected at least {requirements.get('minPages')}")
    if requirements.get("semanticLists") and _num_paragraphs(document) == 0:
        issues.append("semantic Word lists are missing")
    if requirements.get("tables") and not document.tables:
        issues.append("required table is missing")
    if requirements.get("repeatingHeader"):
        if not document.tables or document.tables[0].rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is None:
            issues.append("table header does not repeat")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    letters = [char for char in text if char.isalpha()]
    cyrillic_ratio = (
        sum("\u0400" <= char <= "\u04ff" for char in letters) / len(letters)
        if letters
        else 0.0
    )
    if cyrillic_ratio < 0.5:
        issues.append(f"Cyrillic is not the primary document language ({cyrillic_ratio:.0%})")
    if requirements.get("sources") and not any("source" in paragraph.text.lower() or "\u0418\u0441\u0442\u043e\u0447\u043d" in paragraph.text for paragraph in document.paragraphs):
        issues.append("sources section is missing")
    if requirements.get("assumptions") and not any("\u0414\u043e\u043f\u0443\u0449\u0435\u043d" in paragraph.text for paragraph in document.paragraphs):
        issues.append("assumptions section is missing")
    if source and source_hash and hashlib.sha256(source.read_bytes()).hexdigest() != source_hash:
        issues.append("source file was modified")
    if requirements.get("targetedEdit") and RU["duration_six"] not in text:
        issues.append("targeted edit was not applied")
    if requirements.get("templateFidelity") and "{{" in text:
        issues.append("template placeholders remain")
    critical = [item for item in report.get("issues", []) if item.get("severity") == "critical"]
    if critical:
        issues.append(f"critical QA issues remain: {critical}")
    return issues


def run_matrix(runs: int) -> dict[str, Any]:
    cases = json.loads(CASES_PATH.read_text(encoding="utf-8"))
    results: list[dict[str, Any]] = []
    failures: list[str] = []
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    for case in cases:
        for run_number in range(1, runs + 1):
            run_dir = RESULTS_DIR / case["id"] / f"run-{run_number}"
            _reset_run_dir(run_dir)
            filename = f"{case['id']}.docx"
            output = run_dir / filename
            spec, source = _case_spec(case, filename, run_dir)
            source_hash = hashlib.sha256(source.read_bytes()).hexdigest() if source else None
            spec_path = run_dir / "spec.json"
            spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
            process = subprocess.run([sys.executable, str(BUILDER), str(spec_path), str(output)], capture_output=True, text=True, timeout=180, check=False)
            report_path = Path(f"{output}.artifact-report.json")
            if process.returncode != 0 or not output.exists() or not report_path.exists():
                message = (process.stderr or process.stdout or "builder failed")[-500:]
                failures.append(f"{case['id']} run {run_number}: {message}")
                continue
            report = json.loads(report_path.read_text(encoding="utf-8"))
            pdf = output.with_suffix(".pdf")
            if not pdf.exists():
                failures.append(f"{case['id']} run {run_number}: derived PDF is missing")
                continue
            images = _render_pages(pdf, run_dir)
            if run_number == 1:
                _write_montage(images, run_dir / "montage.png")
            requirement_issues = _assert_requirements(case, output, report, source, source_hash, images)
            failures.extend(f"{case['id']} run {run_number}: {issue}" for issue in requirement_issues)
            results.append(
                {
                    "case": case["id"],
                    "run": run_number,
                    "status": report["status"],
                    "pages": len(images),
                    "imageHashes": [hashlib.sha256(path.read_bytes()).hexdigest() for path in images],
                    "qaChecks": report["qaChecks"],
                }
            )
    expected_cases = [case["id"] for case in cases]
    failures.extend(evaluate_matrix_coverage(results, expected_cases, runs))
    render_digest, evidence_issues = build_render_evidence_digest(results, expected_cases)
    failures.extend(evidence_issues)
    summary = {
        "runsPerCase": runs,
        "caseCount": len(cases),
        "renderEvidence": {
            "algorithm": RENDER_EVIDENCE_ALGORITHM,
            "matrixDigest": render_digest,
        },
        "results": results,
        "failures": failures,
        "passed": not failures,
    }
    (RESULTS_DIR / "matrix-summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--runs", type=int, default=3)
    args = parser.parse_args()
    summary = run_matrix(args.runs)
    print(json.dumps({"passed": summary["passed"], "cases": summary["caseCount"], "runsPerCase": summary["runsPerCase"], "failures": summary["failures"]}, ensure_ascii=False, indent=2))
    return 0 if summary["passed"] else 1


MODULE_SPEC = importlib.util.spec_from_file_location("docx_builder_for_goldens", BUILDER)
assert MODULE_SPEC and MODULE_SPEC.loader
BUILDER_MODULE = importlib.util.module_from_spec(MODULE_SPEC)
MODULE_SPEC.loader.exec_module(BUILDER_MODULE)


if __name__ == "__main__":
    raise SystemExit(main())
